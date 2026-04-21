from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import converter  # Import the converter module
import re
from tkinter import messagebox
from docx.oxml.ns import qn

unicode_fonts = ["Vrinda",
"Vrinda (Headings CS)",
"Shonar Bangla",
"Nirmala UI",
"Kalpurush",
"SolaimanLipi",
"Siyam Rupali",
"Nikosh",
"Noto Sans Bengali",
"Noto Serif Bengali",
"Hind Siliguri",
"Baloo Da",
"Mina",
"Mitra Unicode"]

runs = []

def contains_unicode(text):
    # Check if the text contains Unicode characters
    return bool(re.search(r'[^\x00-\x7F]', text))
unic1 = r"\U+09BC"
minor_replacement_map = {
"Zowr":"Zwor",
" ‡":" †",
"K¨y":"Kz¨",
"MW"+unic1+"‡": "M‡o",
" ‰":" ˆ",
"¯—":"¯Í",
"š—":"šÍ",
"M¬":"Mø",
"Zª":"Î",
" ˆ":" ˆ",
"nÖ":"n«",
"U¨y":"Uz¨",
"Ê":"Ð",
"RvqwK":"RvwqK",
"¯'vqw":"¯’vwq",
"KyÐ":"KzÐ",
"¯c":"¯ú",
"iª&g":"g©",
"aÖy":"aªæ",
"b¨~":"b~¨",
"¶y":"ÿz",
"AÖ¨v":"A¨v",
"Üy":"Üz",
"¯'~":"¯’‚",
"l&µ":"®Œ",
"ù~":"ù‚",
"gœ":"¤œ",
"¤ø¬":"¤ø",
"¯ø¬":"¯ø",
"dÖ":"d«",
"gÖ":"¤ª",
" ‰":" ˆ",
"mÖ":"¯ª",
"iª&w`":"w`©",
"±Ö":"±ª",
"¯Íy":"¯‘",
"iÖ¨":"i¨",
"Zowr":"Zwor",
"å~":"åƒ",
"™¢~":"™¢‚",
"¯‹y":"¯‹z",
"AÖ¨":"A¨",
"m&K":"¯‹",
"¯ø¬¨":"¯ø¨",
"K~":"K‚",
"P~":"P‚",
"Q~":"Q‚",
"S~":"S‚",
"U~":"U‚",
"V~":"V‚",
"W~":"W‚",
"X~":"X‚",
"Z~":"Z‚",
"d~":"d‚",
"eyE":"e~",
"f~":"f‚",
"m¨y":"my¨",
"`¨y":"`yy¨",
"P¨y":"Pz¨",
"k¨y":"ï¨",
"b¨~":"b~¨",
"Z¨y":"Zz¨",
"e¨y":"ey¨",
"RowZ":"RwoZ",
"O&¶":"•ÿ",
"Ö":"ª", 
"Ky":"Kz",
"Py":"Pz",
"Qy":"Qz",
"Sy":"Sz",
"¼y":"¼z",
"Uy":"Uz",
"Vy":"Vz",
"Wy":"Wz",
"O&¸":"½y",
"Xy":"Xz",
"Zy":"Zz",
"dy":"dz",
"fy":"fz",
"oy":"o–",
"‡o":"o‡",
"qª¨v":"q¨v",
"qww":"wqw",
"cowqv":"cwoqv",
"qÖ¨v":"q¨v",
"iÖ¨v":"i¨v",
"i‌ª¨v":"i¨v",
"¤ø­x":"¤øx",
"owr":"wor"
}
def minor_fixes(match):
    matched_char = match.group(0)
    return minor_replacement_map.get(matched_char, matched_char)
pattern = r"|".join(re.escape(key) for key in minor_replacement_map.keys())

def set_run_font(run):
    # 1. Set the primary font name (font.name)
    run.font.name = "SutonnyMJ"
    
    # 2. Force the font name into the XML for all ranges
    # This is often the fix for complex/non-Latin scripts
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "SutonnyMJ")        # Latin/ASCII characters
    rFonts.set(qn("w:hAnsi"), "SutonnyMJ")       # High ANSI (used for Windows characters)
    rFonts.set(qn("w:cs"), "SutonnyMJ")          # Complex Script/Bidi (often for Bengali)

def replace_and_highlight(doc_path, save_path, h_value):
    doc = Document(doc_path)
    doc.save(doc_path+".bkup")
    
    # Create an instance of the Unicode class from converter.py
    unicode_converter = converter.Unicode()
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            print(run)
            if contains_unicode(run.text):
                runs.append(run)
                if run.text =='?' and run.font.name in unicode_fonts:
                    set_run_font(run)
                if contains_unicode(run.text) and run.font.name in unicode_fonts:
            # Convert the text from Unicode to Bijoy
                    converted_text = unicode_converter.convertUnicodeToBijoy(run.text)
                    converted_text = re.sub(pattern, minor_fixes, converted_text)
                    if converted_text[0] == '‡':
                        # Create a new string: '†' + everything from the 2nd character onward
                        converted_text = '†' + converted_text[1:]
                    if converted_text[0] == '‰':
                        # Create a new string: '†' + everything from the 2nd character onward
                        converted_text = 'ˆ' + converted_text[1:]
                    if run.text != converted_text:
                        run.text = converted_text
                        if h_value == 1:
                            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE  # Teal highlight
                        set_run_font(run)  # Set font to Bijoy
                        runs.append(run)
                
                        

    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text =='?' and run.font.name in unicode_fonts:
                            set_run_font(run)
                        if contains_unicode(run.text):
                            if contains_unicode(run.text) and run.font.name in unicode_fonts:
                        # Convert the text from Unicode to Bijoy
                                converted_text = unicode_converter.convertUnicodeToBijoy(run.text)
                                converted_text = re.sub(pattern, minor_fixes, converted_text)
                                if converted_text[0] == '‡':
                        # Create a new string: '†' + everything from the 2nd character onward
                                    converted_text = '†' + converted_text[1:]
                                if converted_text[0] == '‰':
                        # Create a new string: '†' + everything from the 2nd character onward
                                    converted_text = 'ˆ' + converted_text[1:]
                                if run.text != converted_text:
                                    run.text = converted_text
                                    if h_value == 1:
                                        run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE  # Teal highlight
                                    set_run_font(run) # Set font to Bijoy
                                    runs.append(run)
                                

    # Save the document
    doc.save(save_path)
    messagebox.showinfo("Conversion Complete", f"Total words converted: {len(runs)}")