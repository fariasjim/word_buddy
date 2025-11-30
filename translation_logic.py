from docx import Document as doc
from docxlatex import Document as docxlatex
from tkinter import filedialog
from deep_translator import GoogleTranslator
from docx.oxml.ns import nsdecls
from lxml import etree
import time
import customtkinter as ctk
from tkinter import messagebox

M_NAMESPACE = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'

def translator(raw_text):
    translated_text = GoogleTranslator(source='bn', target='en').translate(raw_text)
    return translated_text

def get_simple_omath_text(omath_element):
    if omath_element is None:
        return ""
    text_parts = []
    for text in omath_element.itertext():
        text_parts.append(text)
    return "".join(text_parts).strip()

def main(question_type, file_path, save_path):
    global doc
    doc = doc(file_path)
    if question_type == "mcq":
        step = len(doc.tables)/100
        progress = 0
        for table in doc.tables:
            for r_count, row in enumerate(table.rows):
                if r_count > 0 and r_count < 7:
                    for c_count, cell in enumerate(row.cells):
                        if c_count == 0:
                            next_cell = row.cells[1]
                            if len(next_cell.text) ==0:
                                cell_text = ""
                                iter_cell_text = []
                                for para in cell.paragraphs:
                                    for ele in para._element:
                                        if ele.tag == f"{M_NAMESPACE}oMath":
                                            omath = get_simple_omath_text(ele)
                                            iter_cell_text.append(str(f"#{omath}#"))
                                        else:
                                            if ele.text is not None:
                                                iter_cell_text.append(str(ele.text))
                                    iter_cell_text.append("\n")
                                cell_text = " ".join(iter_cell_text)
                                try:
                                    next_cell.text = translator(cell_text)
                                except Exception as e:
                                    messagebox.showerror("Error", str(e))
                                try:
                                    for para in next_cell.paragraphs:
                                        for run in para.runs:
                                            font = run.font
                                            font.name = "Times New Roman"
                                except:
                                    pass
            progress += step
    elif question_type == "saq":
        for table in doc.tables:
            for r_count, row in enumerate(table.rows):
                if r_count > 1:
                    for c_count, cell in enumerate(row.cells):
                        if c_count == 0:
                            next_cell = row.cells[1]
                            if len(next_cell.text) ==0:
                                cell_text = ""
                                iter_cell_text = []
                                for para in cell.paragraphs:
                                    for ele in para._element:
                                        if ele.tag == f"{M_NAMESPACE}oMath":
                                            omath = get_simple_omath_text(ele)
                                            iter_cell_text.append(str(f"#{omath}#"))
                                        else:
                                            if ele.text is not None:
                                                iter_cell_text.append(str(ele.text))
                                    iter_cell_text.append("\n")
                                cell_text = " ".join(iter_cell_text)
                                try:
                                    next_cell.text = translator(cell_text)
                                except Exception as e:
                                    messagebox.showerror("Error", str(e))
    else:
        pass
    doc.save(save_path)


if __name__ == "__main__":    
    file = filedialog.askopenfilename()
    save = filedialog.asksaveasfilename()
    main(question_type="mcq", file_path=file, save_path=save)
    

    
    

        



                






    



